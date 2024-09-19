import os

import pymed
from dotenv import load_dotenv
from pymed.article import PubMedArticle
from docx import Document

from qwen_agent.agents import Assistant, ArticleAgent
from qwen_agent.agents.assistant import KNOWLEDGE_SNIPPET, KNOWLEDGE_TEMPLATE
from qwen_agent.gui import WebUI
from qwen_agent.llm import get_chat_model
from qwen_agent.llm.schema import Message, USER
from qwen_agent.tools import WebSearch
from qwen_agent.tools.pubmed_search import PubmedSearch
from qwen_agent.tools.simple_doc_parser import parse_pdf
from qwen_agent.utils.utils import list_files
import json
import os

import pandas as pd
import requests_cache
import requests


def app_test():
    os.environ["KMP_DUPLICATE_LIB_OK"] = "true"

    web_search = WebSearch()
    content_items = web_search.call({"query": "pediatric stroke cohort"})
    content_items += web_search.call({"query": "儿童卒中队列的建设"})
    content_items += web_search.call({"query": "youth stroke cohort"})
    content_items += web_search.call({"query": "青年卒中队列的建设"})
    "what is the current status and research progress of clinical research on pediatric stroke"
    "the review of current status and research progress of clinical research on pediatric stroke"
    pubmed_search = PubmedSearch()
    content_items += pubmed_search.call({"query": "pediatric stroke cohort"})
    content_items += pubmed_search.call({"query": "youth stroke cohort"})

    files = [item.file for item in content_items]
    bot = ArticleAgent(llm={"generate_cfg": {"max_input_tokens": 30500}}, files=files)
    messages = [{'role': 'user', 'content': '国内外儿童卒中/青年卒中队列的建设情况'}]
    for rsp in bot.run(messages, full_article=True):
        print(rsp)


PROMPT_TEMPLATE_EN = """You are a writing assistant, please follow the reference paper and extract the content that is useful to me as material based on the given objective.

# Reference paper:
{ref_doc}

# Objective:
{objective}

Please start writing directly, output only the extracted content, do not repeat the objective, do not say irrelevant words, and ensure that the extracted content and the objective remain consistent.

The extracted content should be in JSON format and in English: 
{{
"background": "study background",
"abstract": "study abstract",
"conclusions": "study conclusions",
"importance": "why the study is important?",
"reflection_on_the_objective": "think about the relationship to the objective"
}}
"""

PROMPT_TEMPLATE2_EN = """You are a writing assistant, please follow the reference paper and answer my questions.

# Reference paper:
{ref_doc}

请回答我以下问题：
这篇文章的成果来源于哪些队列研究或干预研究？请列出每个研究的名字，研究类型(队列研究还是干预研究)，研究对象，样本量，团队介绍，研究方法，研究成果，研究的重大意义
"""

pubmed = pymed.PubMed()


def query_pubmed_meta(pmid):
    results = list(pubmed.query("PMID:" + str(pmid), 1))
    if len(results) == 0:
        return None

    article: PubMedArticle = results[0]

    if len(article.authors) > 0:
        first_author = article.authors[0]['firstname'] + " " + article.authors[0]['lastname']
    else:
        first_author = ''

    return {
        'title': article.title,
        'journal': article.journal,
        'first_author': first_author,
        'date_of_publication': str(article.publication_date),
        'abstract': article.abstract,
        'results': article.results or '',
        'conclusions': article.conclusions or '',
        'methods': article.methods or ''
    }


def get_pubmed_meta2(pmid, llm):
    print(f"processing {pmid}")
    json_file = os.path.join(f"f:\\resources\\stroke_text\\{pmid}.meta2.json")

    pdf_file = os.path.join(f"f:\\resources\\stroke\\{pmid}.pdf")
    if not os.path.exists(pdf_file):
        return None

    if not os.path.exists(json_file) or os.path.getsize(json_file) < 100:
        meta = query_pubmed_meta(pmid)

        lines = []
        contents = parse_pdf(pdf_file)
        for content in contents:
            lines.extend([c['text'].strip(" \n") for c in content['content'] if 'text' in c])

        raw_text = "\n".join(lines)
        lines = raw_text.split("\n")
        raw_text = "\n".join([line for line in lines if len(line) >= 5])
        raw_text = raw_text[:20000]

#         prompt1 = """You are a writing assistant, please follow the reference paper and answer my questions.
#
# # Reference paper:
# {ref_doc}
#
# 请提取以下信息：
# 研究对象（object）:
# 样本量（samples）:
# 团队介绍（researchers）:
# 研究背景（background）:
# 研究方法（methods）:
# 研究成果（results）:
# 研究结论（conclusions):
# 研究的意义（importance）:
# """.format(ref_doc=raw_text)
#
#         *_, last = llm.chat(messages=[Message(USER, prompt1)])
#
#         meta['summary'] = last[0].content
#
#         prompt2 = """You are a writing assistant, please follow the reference paper and answer my questions.
#
# # Reference paper:
# {ref_doc}
#
# 请回答我以下问题：
# 这篇文章的成果来源于哪些队列研究或干预研究？请列出每个研究的名字，研究类型(队列研究还是干预研究)，研究对象，样本量，团队介绍，研究方法，研究成果，研究的重大意义
# """.format(ref_doc=raw_text)

        prompt2 = """You are a writing assistant, please summarize the reference paper.

        # Reference paper:
        {ref_doc}

        summary:""".format(ref_doc=raw_text)

        *_, last = llm.chat(messages=[Message(USER, prompt2)])
        meta['summary'] = last[0].content

        with open(json_file, "w", encoding='utf-8') as fp:
            fp.write(json.dumps(meta))

    meta = json.load(open(json_file, "r", encoding='utf-8'))
    return meta

def get_pubmed_meta(pmid, llm):
    print(f"processing {pmid}")
    json_file = os.path.join(f"f:\\resources\\stroke_text\\{pmid}.json")

    pdf_file = os.path.join(f"f:\\resources\\stroke\\{pmid}.pdf")
    if not os.path.exists(pdf_file):
        return None

    if not os.path.exists(json_file) or os.path.getsize(json_file) < 100:
        meta = query_pubmed_meta(pmid)

        lines = []
        contents = parse_pdf(pdf_file)
        for content in contents:
            lines.extend([c['text'].strip(" \n") for c in content['content'] if 'text' in c])

        raw_text = "\n".join(lines)
        lines = raw_text.split("\n")
        raw_text = "\n".join([line for line in lines if len(line) >= 5])
        raw_text = raw_text[:20000]

        objective = "儿童卒中临床研究现状及研究进展综述"
        prompt = PROMPT_TEMPLATE_EN.format(ref_doc=raw_text, objective=objective)
        *_, last = llm.chat(messages=[Message(USER, prompt)])

        parsed_meta = json.loads(last[0].content)
        parsed_meta = {
            'title': meta['title'],
            'journal': meta['journal'],
            'first_author': meta['first_author'],
            'date_of_publication': meta['date_of_publication'],
            **parsed_meta}

        with open(json_file, "w", encoding='utf-8') as fp:
            fp.write(json.dumps(parsed_meta))
    else:
        parsed_meta = json.load(open(json_file, 'r', encoding='utf8'))

    content = "\n".join([f"{key}: {value}" for key, value in parsed_meta.items() if key in ['title', 'journal', 'first_author', 'date_of_publication', 'abstract', 'conclusions', 'importance', 'reflection_on_the_objective']])
    return {'source': f'PMID: {pmid}', 'content': content}


def app_gui():
    os.environ["KMP_DUPLICATE_LIB_OK"] = "true"
    llm = get_chat_model({
            'model': 'qwen2-72b-instruct',
            'generate_cfg': {
                'max_input_tokens': 31000
            }
        })

    requests_cache.install_cache()

    df = pd.read_excel('d:\\all.xlsx')
    df.sort_values(by='pubmed_pubdate', ascending=False, inplace=True)
    df = df[df['impact_factor'] >= 10]

    doc = Document()
    snippets = []
    for index, record in df.iterrows():
        record = record.to_dict()

        try:
            meta = get_pubmed_meta2(int(record['pmid']), llm)
            # doc.add_heading(f"PMID: {record['pmid']} {meta['title']}", level=1)
            # doc.add_paragraph(f"发布时间: {meta['date_of_publication']}")
            # doc.add_paragraph(f"发布期刊: {meta['journal']}")
            # doc.add_paragraph(f"第一作者: {meta['first_author']}")
            #
            # doc.add_paragraph(meta['summary'])
            # doc.add_paragraph('')
            # doc.add_paragraph(meta['cohorts'])
            # doc.save("d:\\all.docx")

            content = f"Title: {meta['title']}\nDate of Publication: {meta['date_of_publication']}\nJournal: {meta['journal']}\nAuthor: {meta['first_author']}\n{meta['summary']}"
            snippets.append(KNOWLEDGE_SNIPPET['en'].format(source=f'{meta["first_author"]}, {meta["date_of_publication"][:4]}', content=content))
        except Exception as e:
            pass

        continue

        # snippet = get_pubmed_meta(record['pmid'], llm)
        # if snippet is None:
        #     continue
        # snippets.append(KNOWLEDGE_SNIPPET['zh'].format(source=snippet['source'], content=snippet['content']))

    knowledge_prompt = '\n\n'.join(snippets)

    # Define the agent
    agent = ArticleAgent(
        llm={
            'model': 'qwen2-72b-instruct',
            'generate_cfg': {
                'max_input_tokens': 100000
            }
        },
        rag_cfg={
            'max_ref_token': 20000,
            'parser_page_size': 250
        },
        name='写作助手',
        system_message="You are helpful assistant to write paper for me",
        description='论文，发言稿，文书写作',
        files=[],
        knowledge=knowledge_prompt
    )

    chatbot_config = {
        'prompt.suggestions': [
            {
                'text': '儿童卒中临床研究现状及研究进展综述'
            }
        ]
    }
    WebUI(agent, chatbot_config=chatbot_config).run()


if __name__ == '__main__':
    # app_test()
    app_gui()
